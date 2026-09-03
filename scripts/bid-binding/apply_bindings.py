from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document


PLACEHOLDER = re.compile(r"\{\{([^{}]+)\}\}")


def normalized(value: str) -> str:
    return re.sub(r"[\s/（）()]", "", value or "")


def replace_in_paragraph(paragraph, values: dict[str, str]) -> None:
    runs = list(paragraph.runs)
    if not runs:
        return
    combined = "".join(run.text for run in runs)
    matches = [match for match in PLACEHOLDER.finditer(combined) if match.group(1) in values]
    for match in reversed(matches):
        start, end = match.span()
        cursor = 0
        start_index = end_index = None
        start_offset = end_offset = 0
        for index, run in enumerate(runs):
            next_cursor = cursor + len(run.text)
            if start_index is None and start < next_cursor:
                start_index = index
                start_offset = start - cursor
            if end <= next_cursor:
                end_index = index
                end_offset = end - cursor
                break
            cursor = next_cursor
        if start_index is None or end_index is None:
            raise ValueError(f"无法定位占位符：{match.group(0)}")
        replacement = str(values[match.group(1)])
        if start_index == end_index:
            text = runs[start_index].text
            runs[start_index].text = text[:start_offset] + replacement + text[end_offset:]
        else:
            prefix = runs[start_index].text[:start_offset]
            suffix = runs[end_index].text[end_offset:]
            runs[start_index].text = prefix + replacement
            for index in range(start_index + 1, end_index):
                runs[index].text = ""
            runs[end_index].text = suffix


def replace_in_cell(cell, values: dict[str, str]) -> None:
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, values)
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                replace_in_cell(nested_cell, values)


def cell_text(cell) -> str:
    return "".join(paragraph.text for paragraph in cell.paragraphs).strip()


def iter_all_paragraphs(document: Document) -> Iterable:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    seen = set()
    for section in document.sections:
        for container in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
            marker = id(container._element)
            if marker in seen:
                continue
            seen.add(marker)
            yield from container.paragraphs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def find_repeated_rows(document: Document, required_placeholders: set[str]):
    matches = []
    for table in document.tables:
        for row in table.rows:
            texts = [cell_text(cell) for cell in row.cells]
            row_placeholders = {
                found.group(1)
                for value in texts
                for found in PLACEHOLDER.finditer(value)
            }
            if required_placeholders <= row_placeholders:
                matches.append((row, texts))
    return matches


def fill_repeated_rows(document: Document, bindings: dict) -> None:
    action_rows = find_repeated_rows(document, {"适用范围", "服务频次"})
    action_bindings = bindings.get("actionRows", [])
    if len(action_rows) != len(action_bindings):
        raise ValueError(f"动作行数量不一致：模板{len(action_rows)}行，映射{len(action_bindings)}行")
    for index, ((row, texts), item) in enumerate(zip(action_rows, action_bindings), start=1):
        if normalized(texts[0]) != normalized(item["expectedTitle"]):
            raise ValueError(
                f"第{index}个动作标题不一致：模板“{texts[0]}”，映射“{item['expectedTitle']}”（{item['id']}）"
            )
        if not item.get("enabled", True):
            row._element.getparent().remove(row._element)
            continue
        for cell in row.cells:
            replace_in_cell(cell, {"适用范围": item["scope"], "服务频次": item["frequency"]})

    staffing_rows = find_repeated_rows(document, {"配置依据", "配置标准", "配置人数"})
    staffing_bindings = bindings.get("staffingRows", [])
    if len(staffing_rows) != len(staffing_bindings):
        raise ValueError(f"岗位行数量不一致：模板{len(staffing_rows)}行，映射{len(staffing_bindings)}行")
    for index, ((row, texts), item) in enumerate(zip(staffing_rows, staffing_bindings), start=1):
        if normalized(texts[0]) != normalized(item["expectedTitle"]):
            raise ValueError(
                f"第{index}个岗位标题不一致：模板“{texts[0]}”，映射“{item['expectedTitle']}”（{item['id']}）"
            )
        if not item.get("enabled", True):
            row._element.getparent().remove(row._element)
            continue
        values = {"配置依据": item["basis"], "配置标准": item["standard"], "配置人数": item["headcount"]}
        for cell in row.cells:
            replace_in_cell(cell, values)


def apply_bindings(template: Path, output: Path, bindings: dict) -> None:
    document = Document(template)
    fill_repeated_rows(document, bindings)
    named = {str(key): str(value) for key, value in bindings.get("named", {}).items()}
    for paragraph in iter_all_paragraphs(document):
        replace_in_paragraph(paragraph, named)
    unresolved = sorted({match.group(0) for paragraph in iter_all_paragraphs(document) for match in PLACEHOLDER.finditer(paragraph.text)})
    if unresolved:
        raise ValueError(f"仍有未填写占位符：{', '.join(unresolved)}")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--bindings", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    apply_bindings(args.template, args.output, json.loads(args.bindings.read_text(encoding="utf-8")))
    print(args.output)


if __name__ == "__main__":
    main()
