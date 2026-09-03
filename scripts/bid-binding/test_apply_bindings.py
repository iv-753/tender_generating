from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from apply_bindings import apply_bindings


ROOT = Path(__file__).resolve().parents[3]
TEMPLATE = ROOT / "output" / "bid-template" / "安序物业_住宅物业服务投标文件_双括号动态母版_清理版.docx"


def cell_text(cell) -> str:
    return "".join(paragraph.text for paragraph in cell.paragraphs).strip()


class ApplyBindingsTest(unittest.TestCase):
    def test_fills_each_repeated_row_from_its_own_binding(self) -> None:
        template = Document(TEMPLATE)
        action_titles = []
        staffing_titles = []
        named_keys = set()
        for table in template.tables:
            for row in table.rows:
                values = [cell_text(cell) for cell in row.cells]
                if "{{适用范围}}" in values and "{{服务频次}}" in values:
                    action_titles.append(values[0])
                if "{{配置依据}}" in values and "{{配置标准}}" in values and "{{配置人数}}" in values:
                    staffing_titles.append(values[0])
        for paragraph in template.paragraphs:
            text = paragraph.text
            while "{{" in text and "}}" in text:
                start = text.index("{{")
                end = text.index("}}", start) + 2
                named_keys.add(text[start + 2:end - 2])
                text = text[end:]
        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    while "{{" in text and "}}" in text:
                        start = text.index("{{")
                        end = text.index("}}", start) + 2
                        named_keys.add(text[start + 2:end - 2])
                        text = text[end:]
        named_keys -= {"适用范围", "服务频次", "配置依据", "配置标准", "配置人数"}

        bindings = {
            "named": {key: f"值-{key}" for key in named_keys},
            "actionRows": [
                {"id": f"action-{index}", "expectedTitle": title, "scope": f"范围-{index}", "frequency": f"频次-{index}"}
                for index, title in enumerate(action_titles, start=1)
            ],
            "staffingRows": [
                {"id": f"staff-{index}", "expectedTitle": title, "basis": f"依据-{index}", "standard": f"标准-{index}", "headcount": str(index)}
                for index, title in enumerate(staffing_titles, start=1)
            ],
        }

        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "result.docx"
            apply_bindings(TEMPLATE, output, bindings)
            generated = Document(output)
            all_text = "\n".join(
                [paragraph.text for paragraph in generated.paragraphs]
                + [cell.text for table in generated.tables for row in table.rows for cell in row.cells]
            )
            self.assertNotIn("{{", all_text)
            self.assertIn("范围-1", all_text)
            self.assertIn(f"频次-{len(action_titles)}", all_text)
            self.assertIn(f"依据-{len(staffing_titles)}", all_text)
            self.assertEqual(len(action_titles), 109)
            self.assertEqual(len(staffing_titles), 6)

    def test_rejects_a_row_title_mismatch(self) -> None:
        template = Document(TEMPLATE)
        action_titles = []
        staffing_titles = []
        for table in template.tables:
            for row in table.rows:
                values = [cell_text(cell) for cell in row.cells]
                if "{{适用范围}}" in values and "{{服务频次}}" in values:
                    action_titles.append(values[0])
                if "{{配置依据}}" in values and "{{配置标准}}" in values and "{{配置人数}}" in values:
                    staffing_titles.append(values[0])
        bindings = {
            "named": {},
            "actionRows": [
                {"id": f"action-{index}", "expectedTitle": "错误动作" if index == 1 else title, "scope": "范围", "frequency": "频次"}
                for index, title in enumerate(action_titles, start=1)
            ],
            "staffingRows": [
                {"id": f"staff-{index}", "expectedTitle": title, "basis": "依据", "standard": "标准", "headcount": "1"}
                for index, title in enumerate(staffing_titles, start=1)
            ],
        }
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "result.docx"
            with self.assertRaisesRegex(ValueError, "动作标题不一致"):
                apply_bindings(TEMPLATE, output, bindings)


if __name__ == "__main__":
    unittest.main()
